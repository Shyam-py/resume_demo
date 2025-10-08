import streamlit as st
from openai import OpenAI
import pdfplumber
import docx2txt
import requests
from bs4 import BeautifulSoup
from io import BytesIO
from docx import Document
import tempfile
import os
import json
import re

# --- CONFIG ---
st.set_page_config(page_title="TEST", layout="wide")
st.title("Resume Optimizer (Testing)")
st.caption("Upload your resume and a job description — get a truthful, ATS-optimized version instantly.")

# ======== 🔑 OPENAI API KEY =========
# Default OpenAI client and model
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
MODEL = "gpt-4o-mini"
TEMPERATURE = 0.1
MAX_OUTPUT_TOKENS = 1600

# --- HELPER FUNCTIONS ---
def extract_text_from_pdf(file_bytes):
    text = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text.append(t)
    return "\n".join(text)

def extract_text_from_docx(file_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        text = docx2txt.process(tmp_path)
    finally:
        os.remove(tmp_path)
    return text or ""

def fetch_text_from_url(url):
    try:
        resp = requests.get(url, timeout=10)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for s in soup(["script", "style", "noscript"]):
            s.decompose()
        main = soup.find("main")
        if main:
            text = main.get_text(separator="\n")
        else:
            divs = soup.find_all("div")
            largest = max(divs, key=lambda d: len(d.get_text())) if divs else soup
            text = largest.get_text(separator="\n")
        return text.strip()
    except Exception as e:
        return f"Error fetching URL: {e}"

def build_system_prompt():
    return (
        "You are a professional resume optimization assistant. "
        "Rewrite resumes to align with job descriptions truthfully and effectively. "
        "Rules:\n"
        "- Never invent roles, dates, or certifications.\n"
        "- Add or rephrase skills from the job posting only if consistent with the candidate’s background.\n"
        "- Format output in clear, ATS-friendly plain text with section headers.\n"
        "- Provide three sections in JSON: optimized_resume, changelog, and suggestions.\n"
        "- The resume must look polished, keyword-rich, and factual."
    )

def build_user_prompt(resume_text, job_text):
    return (
        f"Candidate Resume:\n{resume_text}\n\n"
        f"Job Description:\n{job_text}\n\n"
        "Rewrite the resume to align with the job posting as per the system rules above. "
        "If key skills are missing, list them only in 'suggestions', not inside the resume."
    )

def call_openai_chat(system_prompt, user_prompt):
    try:
        response = client.responses.create(
            model=MODEL,
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=TEMPERATURE,
            max_output_tokens=MAX_OUTPUT_TOKENS,
        )
        return response.output_text, None
    except Exception as e:
        return None, str(e)

def create_docx_from_text(text):
    doc = Document()
    for line in text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        elif line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def parse_json_output(output_text):
    try:
        match = re.search(r"\{[\s\S]*\}", output_text)
        if match:
            return json.loads(match.group(0))
    except Exception:
        pass
    return {
        "optimized_resume": output_text,
        "changelog": "Could not parse changelog from model output.",
        "suggestions": "Could not parse suggestions from model output.",
    }

# --- MAIN UI ---
st.subheader("Upload your resume")
resume_file = st.file_uploader("Upload (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])

st.subheader(" Provide the job description or link")
job_description = st.text_area("Paste job description", height=200)
job_url = st.text_input("Or paste job posting URL")

if st.button("Fetch job posting from URL") and job_url:
    with st.spinner("Fetching job posting..."):
        text = fetch_text_from_url(job_url)
        if text.startswith("Error"):
            st.error(text)
        else:
            job_description = text
            st.success("Fetched job description successfully.")
            st.rerun()


if not resume_file:
    st.info("Please upload a resume to continue.")
    st.stop()

if not job_description.strip():
    st.warning("Please provide a job description or a valid URL.")
    st.stop()

# --- Extract Resume Text ---
with st.spinner("Extracting resume text..."):
    file_bytes = resume_file.read()
    if resume_file.name.endswith(".pdf"):
        resume_text = extract_text_from_pdf(file_bytes)
    elif resume_file.name.endswith(".docx"):
        resume_text = extract_text_from_docx(file_bytes)
    else:
        resume_text = file_bytes.decode(errors="ignore")

# --- Optimize Resume ---
st.subheader("Optimizing your resume...")

system_prompt = build_system_prompt()
user_prompt = build_user_prompt(resume_text, job_description)

with st.spinner("Optimizing resume via OpenAI..."):
    output, err = call_openai_chat(system_prompt, user_prompt)

if err:
    st.error(f"OpenAI API error: {err}")
    st.stop()

parsed_output = parse_json_output(output)
optimized_resume = parsed_output.get("optimized_resume", "")
changelog = parsed_output.get("changelog", "")
suggestions = parsed_output.get("suggestions", "")

# --- Display Results ---
st.success("Resume optimization complete!")

st.markdown("### Optimized Resume (ATS-Friendly)")
st.code(optimized_resume, language="text")

st.markdown("### Changelog")
st.write(changelog)

st.markdown("### Suggestions for Improvement")
st.write(suggestions)

# --- Download Buttons ---
docx_io = create_docx_from_text(optimized_resume)
txt_io = BytesIO(optimized_resume.encode("utf-8"))

col1, col2 = st.columns(2)
with col1:
    st.download_button(
        "Download .docx",
        data=docx_io,
        file_name="optimized_resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
with col2:
    st.download_button(
        "Download .txt",
        data=txt_io,
        file_name="optimized_resume.txt",
        mime="text/plain",
    )


